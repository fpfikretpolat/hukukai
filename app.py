import streamlit as st
import fitz  # PyMuPDF
import os
from google import genai
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
import PIL.Image
from PIL import ImageSequence
import zipfile  # YENİ: UDF (UYAP) arşivlerini açmak için
import re       # YENİ: UDF içindeki etiketleri temizlemek için
import time

# --- SAYFA AYARLARI ---
st.set_page_config(page_title="Hukuk AI Asistanı", page_icon="⚖️", layout="wide")

# GÜNCELLEME: Mobil uyumlu, Sidebar tuşunu bozmayan temiz CSS
gizleme_stili = """
            <style>
            /* Sağ üstteki ana menü (3 nokta) ve Deploy butonunu gizle */
            #MainMenu {visibility: hidden !important;}
            .stDeployButton {display: none !important;} 
            
            /* Alt kısmı ve Made with Streamlit yazısını gizle */
            footer {visibility: hidden !important;}
            
            /* Sağ alttaki geliştirici/yönetici logolarını gizle */
            [data-testid="stToolbar"] {display: none !important;}
            [data-testid="stDecoration"] {display: none !important;}
            .viewerBadge_container {display: none !important;}
            .viewerBadge_link {display: none !important;}
            #st-manage-app-badge {display: none !important;}
            </style>
            """
st.markdown(gizleme_stili, unsafe_allow_html=True)



# --- API KURULUMU ---
load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")

if not api_key:
    st.error("HATA: .env dosyasında GOOGLE_API_KEY bulunamadı! Lütfen kontrol edin.")
    st.stop()

client = genai.Client(api_key=api_key)

# --- UZMAN PROMPT KÜTÜPHANESİ ---
DAVA_PROMPTLARI = {
    "Ceza Hukuku (Ağır Ceza / Asliye)": """
        Sen uzman bir Türk Ceza Avukatısın. Dosyadaki ifadeler, HTS kayıtları, kamera görüntüleme tutanakları ve bilirkişi raporları arasındaki MADDİ ÇELİŞKİLERİ bul.
        Zaman, mekan, silah türü ve kişi teşhislerindeki uyumsuzluklara odaklan. 
        Analizini 3 başlıkta yap: 1. Olay Özeti, 2. Tespit Edilen Çelişkiler, 3. Hukuki Risk ve Savunma Stratejisi (Beraat/İndirim sebepleri).
    """,
    "Aile Hukuku (Boşanma / Velayet)": """
        Sen uzman bir Türk Aile Hukuku Avukatısın. Boşanma dosyalarındaki dilekçeler, tanık beyanları, mesaj kayıtları ve banka dökümlerini incele.
        Özellikle KUSUR tespiti (sadakatsizlik, şiddet, ekonomik baskı) için tanık beyanlarındaki tutarsızlıkları ve yalanları bul. 
        Analizini 3 başlıkta yap: 1. Uyuşmazlık Özeti, 2. İddia ve Tanık Beyanlarındaki Çelişkiler, 3. Nafaka, Tazminat ve Velayet İçin Risk/Strateji Analizi.
    """,
    "Miras Hukuku (Muris Muvazaası / Tenkis)": """
        Sen uzman bir Türk Miras Hukuku Avukatısın. Dosyadaki tapu senetleri, resmi vasiyetnameler, banka dekontları ve nüfus kayıtlarını çapraz incele.
        Gerçek bedel ile tapuda gösterilen bedel arasındaki uçurumları, mal kaçırma (muris muvazaası) kastını ve saklı pay (tenkis) ihlallerini bul.
        Analizini 3 başlıkta yap: 1. Dava Özeti, 2. Belge ve Bedel Çelişkileri (Muvazaa Tespiti), 3. Hukuki Risk ve İspat Stratejisi.
    """,
    "Borçlar ve Ticaret Hukuku (Alacak / Sözleşme)": """
        Sen uzman bir Türk Ticaret ve Borçlar Hukuku Avukatısın. Ticari defterler, faturalar, sözleşme maddeleri, sevk irsaliyeleri ve ihtarnameleri incele.
        Temerrüt tarihlerindeki hataları, sözleşmeye aykırılıkları, imza inkarını ve faiz başlangıç tarihlerindeki tutarsızlıkları tespit et.
        Analizini 3 başlıkta yap: 1. Sözleşme/İhtilaf Özeti, 2. Belge ve Beyan Çelişkileri, 3. Zamanaşımı/Hak Düşürücü Süre Riskleri ve Dava Stratejisi.
    """,
    "İcra ve İflas Hukuku (İlamsız/İlamlı Takip)": """
        Sen uzman bir Türk İcra ve İflas Hukuku Avukatısın. Dosyadaki takip taleplerini, ödeme/icra emirlerini, tebligat mazbatalarını, itiraz dilekçelerini ve hesap özetlerini çok dikkatli incele.
        Özellikle şu usul hatalarını ve çelişkileri ara: Tebligat Kanunu'na aykırı (usulsüz) tebligat yapılıp yapılmadığı, hak düşürücü itiraz sürelerinin kaçırılıp kaçırılmadığı, takip talebi ile ödeme emri arasındaki bedel/faiz uyuşmazlıkları ve yetkisiz icra dairesi seçimi.
        Analizini 3 başlıkta yap: 1. İcra Dosyası Özeti, 2. Tespit Edilen Usul Hataları ve Çelişkiler (Tebligat, Süre, Faiz Oranları), 3. Hukuki Risk ve Şikayet/İtiraz Stratejisi.
    """,
    "Eğitim Hukuku (Üniversiteden İlişik Kesme / Disiplin)": """
        Sen uzman bir Türk İdare ve Eğitim Hukuku Avukatısın. Dosyadaki üniversite senatosu/yönetim kurulu kararlarını, disiplin soruşturması raporlarını, öğrenci savunma tutanaklarını ve transkript belgelerini çok dikkatli incele.
        Özellikle şu USUL ve ESAS hatalarını ara: 
        1. 2547 Sayılı YÖK Kanunu (md. 44 ve 54) ve Öğrenci Disiplin Yönetmeliği'ne aykırılıklar.
        2. Öğrenciye usulüne uygun savunma hakkı (en az 7 gün süre) verilip verilmediği.
        3. Disiplin soruşturmasının kanuni süreler içinde açılıp sonuçlandırılıp sonuçlandırılmadığı (Zamanaşımı).
        4. "Azami süre" (öğrenim süresinin dolması) nedeniyle kaydı silinen öğrenciler için ek sınav / ek süre haklarının kullandırılıp kullandırılmadığı.
        Analizini 3 başlıkta yap: 1. İlişik Kesme Kararının Özeti, 2. Tespit Edilen Usul ve Yönetmelik İhlalleri (Çok Kritik), 3. İptal Davası ve "Yürütmeyi Durdurma (YD)" Stratejisi.
    """
}

# --- FONKSİYONLAR ---
def dosya_isleyici(yuklenen_dosyalar):
    islenmis_icerik = []
    for dosya in yuklenen_dosyalar:
        dosya_adi = dosya.name.lower()
        islenmis_icerik.append(f"\n\n{'='*50}\n📁 DOSYA KAYNAĞI: {dosya.name}\n{'='*50}\n")
        try:
            # 1. SENARYO: PDF
            if dosya_adi.endswith(".pdf"):
                doc = fitz.open(stream=dosya.read(), filetype="pdf")
                for sayfa in doc:
                    metin = sayfa.get_text().strip()
                    if len(metin) > 20:
                        islenmis_icerik.append(f"\n--- {dosya.name} / SAYFA {sayfa.number + 1} ---\n")
                        islenmis_icerik.append(metin)
                    else:
                        st.toast(f"📸 {dosya.name} - Sayfa {sayfa.number + 1} resim olarak taranıyor...")
                        pix = sayfa.get_pixmap(dpi=150) 
                        img = PIL.Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        islenmis_icerik.append(f"\n--- {dosya.name} / SAYFA {sayfa.number + 1} (GÖRSEL TARAMA) ---\n")
                        islenmis_icerik.append(img)
            
            # 2. SENARYO: UDF (UYAP FORMATI)
            elif dosya_adi.endswith(".udf"):
                st.toast(f"📄 {dosya.name} (UDF) dosyası ayrıştırılıyor...")
                with zipfile.ZipFile(dosya) as archive:
                    if 'content.xml' in archive.namelist():
                        xml_content = archive.read('content.xml').decode('utf-8', errors='ignore')
                        # UYAP metinleri genellikle CDATA içinde yer alır
                        cdata_match = re.search(r'<!\[CDATA\[(.*?)\]\]>', xml_content, re.DOTALL)
                        raw_text = cdata_match.group(1) if cdata_match else xml_content
                        # XML ve HTML etiketlerini tamamen temizle
                        temiz_metin = re.sub(r'<[^>]+>', ' ', raw_text)
                        temiz_metin = re.sub(r'\s+', ' ', temiz_metin).strip()
                        
                        islenmis_icerik.append(f"\n--- {dosya.name} ---\n")
                        islenmis_icerik.append(temiz_metin)
                    else:
                        st.error(f"❌ {dosya.name} geçerli bir UYAP dosyası değil.")

            # 3. SENARYO: TIF, JPG, PNG GÖRÜNTÜLERİ
            elif dosya_adi.endswith((".tif", ".tiff", ".jpg", ".jpeg", ".png")):
                st.toast(f"🖼️ {dosya.name} görüntü dosyası işleniyor...")
                img = PIL.Image.open(dosya)
                if hasattr(img, "n_frames") and img.n_frames > 1:
                    for i, frame in enumerate(ImageSequence.Iterator(img)):
                        islenmis_icerik.append(f"\n--- {dosya.name} / SAYFA {i + 1} ---\n")
                        islenmis_icerik.append(frame.convert("RGB"))
                else:
                    islenmis_icerik.append(f"\n--- {dosya.name} ---\n")
                    islenmis_icerik.append(img.convert("RGB"))
                    
        except Exception as e:
            st.error(f"❌ '{dosya.name}' okuma hatası: {e}")
    return islenmis_icerik

def davayi_analiz_et(gorsel_ve_metin_listesi, secilen_prompt):
    max_deneme = 3
    bekleme_suresi = 3 
    
    for deneme in range(max_deneme):
        try:
            gonderilecek_paket = [secilen_prompt] + gorsel_ve_metin_listesi
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=gonderilecek_paket
            )
            return response.text
            
        except Exception as e:
            hata_mesaji = str(e)
            if "503" in hata_mesaji or "UNAVAILABLE" in hata_mesaji or "429" in hata_mesaji:
                if deneme < max_deneme - 1:
                    st.toast(f"⚠️ Google sunucuları yoğun. {bekleme_suresi} sn içinde tekrar deneniyor... (Deneme {deneme + 1}/{max_deneme})")
                    time.sleep(bekleme_suresi)
                    bekleme_suresi *= 2 
                    continue 
            return f"❌ AI Analiz hatası: {hata_mesaji}\n\nLütfen 1-2 dakika bekleyip tekrar deneyin."

def word_olustur(rapor_metni):
    doc = Document()
    doc.add_heading('Hukuki Analiz Raporu', 0)
    doc.add_paragraph(rapor_metni)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- HAFIZA (SESSION STATE) AYARI ---
if "analiz_sonucu" not in st.session_state:
    st.session_state.analiz_sonucu = None

# --- ARAYÜZ (UI) BAŞLANGICI ---
st.title("⚖️ Hukuk AI Asistanı")
st.markdown("Hukuk dallarına özel uzmanlaşmış yapay zeka ile dosya analizi.")

st.sidebar.header("⚙️ Analiz Ayarları")
secilen_kategori = st.sidebar.selectbox(
    "Dava Türünü Seçin",
    options=list(DAVA_PROMPTLARI.keys()),
    index=0
)

st.sidebar.markdown("---")
st.sidebar.info(f"🧠 Aktif Uzmanlık: **{secilen_kategori}**")

# GÜNCELLEME: UDF Desteği arayüze eklendi
yuklenen_dosyalar = st.file_uploader(
    "Dosyaları (UDF, PDF, TIF, JPG, PNG) buraya sürükleyin", 
    type=["udf", "pdf", "tif", "tiff", "jpg", "jpeg", "png"], 
    accept_multiple_files=True
)

if yuklenen_dosyalar:
    st.info(f"📂 {len(yuklenen_dosyalar)} dosya eklendi. Analiz Modu: **{secilen_kategori}**")
    
    if st.button("Seçili Branşta Analiz Et", type="primary", use_container_width=True):
        with st.spinner(f"⏳ Yapay zeka {secilen_kategori} uzmanı olarak dosyaları okuyor..."):
            islenmis_icerik = dosya_isleyici(yuklenen_dosyalar)
            if islenmis_icerik:
                aktif_prompt = DAVA_PROMPTLARI[secilen_kategori]
                st.session_state.analiz_sonucu = davayi_analiz_et(islenmis_icerik, aktif_prompt)

if st.session_state.analiz_sonucu:
    if st.session_state.analiz_sonucu.startswith("❌"):
        st.error("Sistem şu anda yanıt veremedi. Lütfen birazdan tekrar deneyin.")
        st.warning(st.session_state.analiz_sonucu) 
    else:
        st.success("✅ Analiz Tamamlandı!")
        st.markdown("---")
        st.markdown(st.session_state.analiz_sonucu)
        st.markdown("---")
        word_dosyasi = word_olustur(st.session_state.analiz_sonucu)
        st.download_button(
            label="📄 Raporu Word Dosyası Olarak İndir",
            data=word_dosyasi,
            file_name="hukuki_analiz_raporu.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
